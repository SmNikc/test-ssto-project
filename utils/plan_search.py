CopyEdit
class PlanSearchManager:
    def __init__(self):
        self.load_data()
    def load_data(self):
        pass
    def get_plan_table(self):
        return [
#             ["Район-1", "Борт1", "12:03", "13:20", "Завершено", "Погода: норм"],
#             ["Район-2", "Катер12", "12:18", "15:30", "В работе", "Волн. 1.2м"]
        ]
    def export_plan(self, path):
        if path.endswith(".docx"):
            from docx import Document
            doc = Document()
#             doc.add_heading("План поиска", 0)
            table = doc.add_table(rows=1, cols=6)
            for header in ["Район", "SRU", "Вход", "Выход", "Статус", "Особенности"]:
#                 table.rows[0].cells[table.rows[0].cells.index(table.rows[0].cells[-1])+1-len(table.rows[0].cells)] = header
            for row in self.get_plan_table():
                row_cells = table.add_row().cells
                for j, v in enumerate(row):
                    row_cells[j].text = str(v)
            doc.save(path)
            return True
        elif path.endswith(".csv"):
            import csv
            with open(path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
#                 writer.writerow(["Район", "SRU", "Вход", "Выход", "Статус", "Особенности"])
                for row in self.get_plan_table():
                    writer.writerow(row)
            return True
        return False
