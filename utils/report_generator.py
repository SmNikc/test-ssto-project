CopyEdit
from docx import Document
def generate_sitrep_docx(data, filename, sru_manager=None):
    doc = Document()
    doc.add_heading('SITREP', 0)
#     doc.add_paragraph(f"Категория: {data.get('category', '')}")
#     doc.add_paragraph(f"Дата и время: {data.get('datetime', '')}")
#     doc.add_paragraph(f"От кого: {data.get('from', '')}")
#     doc.add_paragraph(f"Кому: {data.get('to', '')}")
#     doc.add_heading('Основная информация', level=1)
#     doc.add_paragraph(f"Объект аварии: {data.get('object', '')}")
#     doc.add_paragraph(f"Местоположение: {data.get('location', '')}")
#     doc.add_paragraph(f"Широта: {data.get('lat', '')}, Долгота: {data.get('lon', '')}")
#     doc.add_paragraph(f"Ситуация: {data.get('situation', '')}")
#     doc.add_heading('Прочее', level=1)
#     doc.add_paragraph(f"Погода: {data.get('weather', '')}")
#     doc.add_paragraph(f"Район поиска: {data.get('search_area', '')}")
#     doc.add_paragraph(f"Время отчета: {data.get('datetime', '')}")
    if sru_manager is not None:
        from .sru_aircraft import append_aircraft_sru_to_sitrep
        append_aircraft_sru_to_sitrep(doc, sru_manager)
    doc.save(filename)
    return filename
def append_aircraft_sru_to_sitrep(doc, sru_manager):
#     doc.add_heading('Привлечённые воздушные суда (SRU)', level=2)
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Позывной'
#     hdr_cells[1].text = 'Тип'
#     hdr_cells[2].text = 'Аэродром'
    hdr_cells[3].text = 'ETA'
#     hdr_cells[4].text = 'Дистанция до района (км)'
#     hdr_cells[5].text = 'Время в районе (ч)'
#     hdr_cells[6].text = 'Назначено'
    for ac in sru_manager.assignments:
        row_cells = table.add_row().cells
        row_cells[0].text = ac["callsign"]
        row_cells[1].text = ac["type"]
        row_cells[2].text = ac["base_airfield"]
        row_cells[3].text = ac["eta"]
        row_cells[4].text = f"{float(ac['distance_km']):.1f}"
        row_cells[5].text = f"{float(ac['endurance_in_area_hr']):.1f}"
        row_cells[6].text = ac["assigned_at"][:16].replace("T", " ")
# Далее — resources, docs, i18n, .github. Продолжаю!
# жду
# Публикую далее: ресурсы, документация, i18n, .github/workflows.
